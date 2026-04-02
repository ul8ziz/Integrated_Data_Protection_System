# -*- coding: utf-8 -*-
"""Generates docs/REQUIREMENTS_SECTION_3_2.docx — section 3.2 requirements (Arabic)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt


def _rtl(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def _font_ar(paragraph, size_pt=12):
    for run in paragraph.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii", "Arial")
        run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi", "Arial")
        run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs", "Arial")
        run.font.size = Pt(size_pt)


def _add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _rtl(p)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
    return p


def _add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _rtl(p)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    _font_ar(p)
    return p


def main():
    out = Path(__file__).resolve().parent / "REQUIREMENTS_SECTION_3_2.docx"
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _rtl(title)
    tr = title.add_run("3.2 تحديد المتطلبات")
    tr.bold = True
    tr.font.name = "Arial"
    tr.font.size = Pt(18)

    intro = (
        "هناك عدة كيانات في النظام المقترح، ولكل كيان مجموعة من المتطلبات. "
        "وقد قُسّمت متطلبات النظام إلى أقسام رئيسة كما يلي:"
    )
    _add_para(doc, intro)

    _add_heading(doc, "3.2.1 متطلبات عامة لكل المستخدمين", level=2)
    _add_para(
        doc,
        "يُعد المستخدم من أهم عناصر النظام، وتحقيق متطلباته جزء من تحقيق متطلبات النظام. "
        "وتتلخص متطلبات المستخدم العامة فيما يلي:",
    )
    _add_para(doc, "1. توفير الأمان والحماية لبيانات المستخدمين داخل التطبيق، بما يتوافق مع سياسات المنظمة ومعايير حماية البيانات (مثل GDPR/HIPAA حيث ينطبق).")
    _add_para(doc, "2. تصميم واجهات واضحة وسهلة الاستخدام، مع تغذية راجعة فورية (نجاح/خطأ/تحميل).")
    _add_para(doc, "3. توفير إرشادات وتعليمات حول استخدام التطبيق والوظائف الأساسية (التسجيل، التحليل، المراقبة، لوحة الإدارة).")

    _add_heading(doc, "3.2.2 المتطلبات الوظيفية", level=2)
    _add_para(doc, "المتطلبات الوظيفية هي ما يجب أن يؤديه النظام للمستخدمين، وهي كالتالي:")
    _add_para(doc, "أولاً: متطلبات مسؤول النظام (Admin)", bold=True)
    rows = [
        ("أ", "إنشاء حسابات المستخدمين وإدارتها (بما في ذلك الموافقة على التسجيلات المعلّقة)."),
        ("ب", "تسجيل الدخول والمصادقة."),
        ("ج", "إضافة مستخدمين وتحديث حالاتهم (مثل: معلّق، مفعّل)."),
        ("د", "إدارة الأقسام/التنظيم والسياسات والتنبيهات حسب نطاق النظام."),
        ("هـ", "إدارة النظام ومراقبة التشغيل والتقارير."),
        ("و", "التحكم بالإعدادات (مثل تكامل المراقبة مع MyDLP حيث يُطبَّق)."),
        ("ز", "عرض تقارير البيانات المكتشفة (نتائج التحليل باستخدام Presidio وما يعادله)."),
        ("ح", "عرض تقارير وإحصائيات المستخدمين والنشاط (حسب الصلاحيات)."),
        ("ط", "دور مدير القسم (Manager): إدارة مستخدمي القسم نفسه وفق القيود المطبّقة في الخادم."),
    ]
    for letter, text in rows:
        _add_para(doc, f"{letter}. {text}")

    _add_para(doc, "ثانياً: المتطلبات الوظيفية للمستخدم العادي (User)", bold=True)
    urows = [
        ("1", "إنشاء حساب (يُنشأ عادة بحالة «في انتظار الموافقة»)."),
        ("2", "تسجيل الدخول."),
        ("3", "تحليل النصوص وتحليل الملفات لاكتشاف البيانات الحساسة."),
        ("4", "استخدام وظائف المراقبة/البريد التجريبي حسب ما يوفّره النظام."),
    ]
    for num, text in urows:
        _add_para(doc, f"{num}) {text}")

    _add_heading(doc, "3.2.3 المتطلبات غير الوظيفية (Non-functional requirements)", level=2)
    _add_para(
        doc,
        "في هذا القسم يُذكر لكل متطلب الدليل التقني أو الوسيلة المستخدمة لتحقيقه في التصميم/التنفيذ المقترح، "
        "وليس الاكتفاء بالوصف العام.",
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "المتطلب"
    hdr[1].text = "الوصف"
    hdr[2].text = "أدلة وشواهد تقنية (كيف يُحقَّق)"
    for c in hdr:
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _rtl(p)
            for run in p.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(11)

    nfr_rows = [
        (
            "الموثوقية",
            "الحفاظ على سلامة البيانات واتساقها وتقليل فقدانها عند الأعطال.",
            "MongoDB كمخزن تطبيقي مع Motor و Beanie؛ فهارس على الحقول الحرجة؛ تسجيل تشغيل منظم إلى ملف؛ نسخ احتياطي واستعادة ضمن إجراءات التشغيل (أدوات MongoDB).",
        ),
        (
            "الكفاءة والاستمرارية",
            "أداء مقبول وقدرة على العمل دون توقف غير مبرر.",
            "FastAPI مع Uvicorn؛ حدود حجم الملفات في مسارات الرفع؛ middleware لفحص محتوى الملفات المرفوعة؛ التحقق من اتصال MongoDB؛ عند التوسع: موازنة حمل ونسخ خلفية.",
        ),
        (
            "الأمنية",
            "حماية النظام من الوصول غير المصرّح به والهجمات الشائعة.",
            "OAuth2 + JWT؛ تجزئة كلمات المرور بـ bcrypt (passlib)؛ التحقق من المدخلات بـ Pydantic؛ فصل طبقات API/خدمة/بيانات؛ نشر الواجهة عبر HTTPS (TLS).",
        ),
        (
            "السرية وخصوصية البيانات",
            "ضمان عدم تعرّض البيانات الحساسة إلا للمخولين.",
            "TLS لحماية النقل؛ تخزين كلمات المرور مجزأة فقط؛ RBAC (Admin / Manager / User) مع فحص الصلاحية في الخادم؛ Microsoft Presidio؛ تكامل MyDLP CE للمراقبة حيث ينطبق.",
        ),
        (
            "سهولة الاستخدام",
            "إمكانية الاستخدام من فئات المستخدمين دون عوائق كبيرة.",
            "واجهة HTML/CSS/JavaScript منظمة؛ نماذج بحقول واضحة وحالات تحميل ورسائل خطأ منظمة.",
        ),
        (
            "الإتاحة والتوفر",
            "توفر الخدمة عند الطلب ضمن الاتفاقيات المحددة.",
            "تشغيل الخدمة كخادم ويب متاح؛ middleware للتحقق من اتصال MongoDB؛ تخطيط صيانة ونسخ احتياطي؛ عند الحاجة reverse proxy مثل Nginx.",
        ),
        (
            "المرونة",
            "قابلية التكيّف مع تغيّر المتطلبات أو بيئة العرض.",
            "واجهة ويب متجاوبة؛ هيكلة وحدات FastAPI (routers) قابلة للتوسعة؛ فصل إعدادات التكامل (متغيرات بيئة لـ MongoDB وMyDLP وPresidio وغيرها) عن الكود.",
        ),
        (
            "سهولة الصيانة",
            "تسهيل اكتشاف الأخطاء وإصلاحها من قبل فريق الصيانة المخوّل.",
            "فصل الطبقات (مسارات API، خدمات، نماذج Beanie)؛ توثيق API (OpenAPI/Swagger مع FastAPI)؛ تسجيل منظم دون تسجيل أسرار أو رموز كاملة.",
        ),
        (
            "إدارة الصلاحيات",
            "لوحة تحكم ووظائف تختلف بحسب الدور.",
            "RBAC مبني على دور المستخدم (Admin وManager وRegular)؛ اعتماديات FastAPI (مثل get_current_user وget_current_admin وget_current_admin_or_manager)؛ JWT يتضمن معرف المستخدم والدور؛ منع الاعتماد على صلاحيات الواجهة فقط.",
        ),
    ]

    for name, desc, evidence in nfr_rows:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = desc
        row[2].text = evidence
        for cell in row:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _rtl(p)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)

    _add_para(
        doc,
        "ملاحظة: عمود «أدلة وشواهد تقنية» يربط كل متطلباً بتقنية أو بروتوكول أو ممارسة يمكن الإشارة إليها في التقرير أو الملحق التقني.",
    )

    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
