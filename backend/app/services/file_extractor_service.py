"""
Service for extracting text from various file formats
"""
from typing import Optional
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)

class FileTextExtractor:
    """Service for extracting text from different file formats"""
    
    def __init__(self):
        """Initialize file extractor with available libraries"""
        self.supported_formats = {
            '.txt': self._extract_txt,
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.doc': self._extract_doc,
            '.xlsx': self._extract_xlsx,
            '.xls': self._extract_xls,
        }
    
    def extract_text(self, file_path: str, file_content: Optional[bytes] = None) -> str:
        """
        Extract text from a file
        
        Args:
            file_path: Path to the file (used to determine file type)
            file_content: Optional file content as bytes (if file is already in memory)
            
        Returns:
            Extracted text as string
            
        Raises:
            ValueError: If file format is not supported
            Exception: If extraction fails
        """
        # Get file extension
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported formats: {', '.join(self.supported_formats.keys())}")
        
        try:
            extractor_func = self.supported_formats[file_ext]
            return extractor_func(file_path, file_content)
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise Exception(f"Failed to extract text from file: {str(e)}")
    
    def _extract_txt(self, file_path: str, file_content: Optional[bytes] = None) -> str:
        """Extract text from plain text file"""
        try:
            if file_content:
                # Try different encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252']:
                    try:
                        return file_content.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                # Fallback: decode with errors='ignore'
                return file_content.decode('utf-8', errors='ignore')
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            raise
    
    def _extract_pdf(self, file_path: str, file_content: Optional[bytes] = None) -> str:
        """Extract text from PDF file"""
        text = ""
        
        # Try pdfplumber first (better for complex PDFs)
        try:
            import pdfplumber
            if file_content:
                import io
                pdf_file = io.BytesIO(file_content)
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            else:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            
            if text.strip():
                return text
        except ImportError:
            logger.warning("pdfplumber not available, trying PyPDF2")
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")
        
        # Fallback to PyPDF2
        try:
            import PyPDF2
            if file_content:
                import io
                pdf_file = io.BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
            else:
                pdf_reader = PyPDF2.PdfReader(file_path)
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            return text
        except ImportError:
            raise ImportError("Neither pdfplumber nor PyPDF2 is installed. Please install at least one: pip install pdfplumber or pip install PyPDF2")
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise
    
    def _extract_docx(self, file_path: str, file_content: Optional[bytes] = None) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            import io
            
            if file_content:
                doc_file = io.BytesIO(file_content)
                doc = Document(doc_file)
            else:
                doc = Document(file_path)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("python-docx is not installed. Please install: pip install python-docx")
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise
    
    def _extract_doc(self, file_path: str, file_content: Optional[bytes] = None) -> str:
        """Extract text from DOC file (old Word format)"""
        # DOC format is complex and requires additional libraries
        # For now, we'll raise an error suggesting conversion to DOCX
        raise ValueError(
            "DOC format (old Word format) is not directly supported. "
            "Please convert the file to DOCX format first."
        )
    
    def _extract_xlsx(self, file_path: str, file_content: Optional[bytes] = None) -> str:
        """Extract text from XLSX file"""
        try:
            import openpyxl
            import io
            
            if file_content:
                workbook_file = io.BytesIO(file_content)
                workbook = openpyxl.load_workbook(workbook_file, data_only=True)
            else:
                workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            text_parts = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip():
                        text_parts.append(row_text)
                text_parts.append("")  # Empty line between sheets
            
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("openpyxl is not installed. Please install: pip install openpyxl")
        except Exception as e:
            logger.error(f"Error extracting XLSX: {e}")
            raise
    
    def _extract_xls(self, file_path: str, file_content: Optional[bytes] = None) -> str:
        """Extract text from XLS file (old Excel format)"""
        # XLS format requires xlrd library
        try:
            import xlrd
            import io
            
            if file_content:
                workbook_file = io.BytesIO(file_content)
                workbook = xlrd.open_workbook(file_contents=file_content)
            else:
                workbook = xlrd.open_workbook(file_path)
            
            text_parts = []
            for sheet_name in workbook.sheet_names():
                sheet = workbook.sheet_by_name(sheet_name)
                text_parts.append(f"Sheet: {sheet_name}")
                for row_idx in range(sheet.nrows):
                    row_values = [str(cell.value) if cell.value else "" for cell in sheet.row(row_idx)]
                    row_text = " | ".join(row_values)
                    if row_text.strip():
                        text_parts.append(row_text)
                text_parts.append("")  # Empty line between sheets
            
            return "\n".join(text_parts)
        except ImportError:
            raise ValueError(
                "XLS format (old Excel format) requires xlrd library. "
                "Please install: pip install xlrd, or convert to XLSX format."
            )
        except Exception as e:
            logger.error(f"Error extracting XLS: {e}")
            raise
    
    def is_supported(self, file_path: str) -> bool:
        """Check if file format is supported"""
        file_ext = Path(file_path).suffix.lower()
        return file_ext in self.supported_formats
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats"""
        return list(self.supported_formats.keys())

