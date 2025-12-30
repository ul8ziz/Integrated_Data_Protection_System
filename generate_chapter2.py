import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_rtl(paragraph):
    """Sets the paragraph direction to RTL."""
    pPr = paragraph._p.get_or_add_pPr()
    section = OxmlElement('w:bidi')
    section.set(qn('w:val'), '1')
    pPr.append(section)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_run_rtl(run):
    """Sets the run direction to RTL (for mixed text)."""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    set_rtl(heading)
    for run in heading.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(16 if level == 1 else 14)
        run.font.color.rgb = RGBColor(0, 0, 0) # Black
        set_run_rtl(run)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    set_run_rtl(run)
    return p

def add_bullet_point(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_run_rtl(run)

def create_document():
    doc = Document()

    # Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # --- Chapter 2 Title Page ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("الفصل الثاني")
    run.font.name = 'Arial'
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph() # Spacer
    doc.add_paragraph() # Spacer

    sections = [
        "2.1 خلفية الدراسة (Background)",
        "2.2 الدراسات السابقة (Literature Review)",
        "2.3 النظام المقترح (Proposal System)",
        "2.4 النظرة العامة للنظام (System Overview)",
        "2.5 آلية عمل النظام (System Working Procedure)",
        "2.6 دراسة الجدوى (Feasibility Study)",
        "2.7 إدارة المخاطر (Risks Management)"
    ]

    for section in sections:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
        set_rtl(p)
        run = p.add_run(section)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        set_run_rtl(run)

    doc.add_page_break()

    # --- Content ---

    # 2.1 Background
    add_heading(doc, "2.1 خلفية الدراسة (Background)", level=2)
    text_bg = """يشهد العالم اليوم ثورة رقمية هائلة أدت إلى تزايد غير مسبوق في حجم البيانات التي يتم إنتاجها وتداولها يومياً داخل المؤسسات. ومع هذا التطور، برزت قضية "أمن البيانات" (Data Security) كواحدة من أهم التحديات التي تواجه المنظمات والحكومات على حد سواء. لم يعد الخطر يقتصر فقط على الهجمات الخارجية، بل أصبح "التسريب الداخلي" للبيانات الحساسة (Insider Threats) يشكل تهديداً خطيراً لسمعة المؤسسات ووضعها المالي والقانوني.

تعتمد المؤسسات تقليدياً على أنظمة منع فقدان البيانات (Data Loss Prevention - DLP) لحماية أصولها المعلوماتية. تعمل هذه الأنظمة عادةً بناءً على قواعد ثابتة وتعبيرات نمطية (Regular Expressions) للبحث عن أنماط محددة مثل أرقام الهواتف أو بطاقات الائتمان. ورغم فعالية هذه الطرق في الحالات البسيطة، إلا أنها تعاني من قصور كبير في التعامل مع البيانات غير المهيكلة (Unstructured Data) وتفتقر إلى القدرة على فهم "سياق" النص، مما يؤدي إلى معدلات عالية من الإنذارات الكاذبة (False Positives).

من هنا، ظهرت الحاجة لدمج تقنيات الذكاء الاصطناعي، وتحديداً معالجة اللغات الطبيعية (NLP)، في أنظمة الحماية لتمكينها من فهم السياق واكتشاف البيانات الحساسة بدقة أعلى، وهو ما يسعى هذا المشروع لتحقيقه من خلال نظام "Athier"."""
    for paragraph in text_bg.split('\n\n'):
        add_paragraph(doc, paragraph.strip())
    doc.add_paragraph()

    # 2.2 Literature Review
    add_heading(doc, "2.2 الدراسات السابقة (Literature Review)", level=2)
    add_paragraph(doc, "تمت مراجعة وتحليل مجموعة من المشاريع والأنظمة العالمية المفتوحة المصدر التي تعمل في مجال حماية البيانات (DLP)، وفيما يلي استعراض لأهم ثلاثة مشاريع تتشابه في آليتها مع أهداف مشروعنا:")

    # Study 1: MyDLP
    add_heading(doc, "2.2.1 مشروع (MyDLP Community Edition): مراقبة الشبكة ومنع التسرب", level=3)
    add_paragraph(doc, "يُعد MyDLP واحداً من المشاريع الرائدة والمفتوحة المصدر في مجال منع فقدان البيانات. يركز بشكل أساسي على مراقبة حركة البيانات عبر الشبكة (Network DLP) وحماية نقاط النهاية.")
    
    add_paragraph(doc, "آلية عمل النظام:", bold=True)
    mech_1 = [
        "1. الاعتراض (Interception): يستخدم بروتوكول ICAP لاعتراض حركة مرور الويب والبريد الإلكتروني من الخوادم الوسيطة (Proxy/SMTP Servers).",
        "2. التحليل (Inspection): يقوم بفحص محتوى الملفات والنصوص المارة بحثاً عن بيانات حساسة (مثل أرقام الهويات) باستخدام قواعد محددة مسبقاً.",
        "3. اتخاذ القرار (Action): بناءً على نتيجة الفحص، يقوم النظام إما بالسماح للبيانات بالمرور (Pass)، أو منعها (Block)، أو وضعها في الحجر (Quarantine) للمراجعة.",
        "4. الإدارة المركزية: يتم التحكم في جميع السياسات ومراجعة السجلات عبر واجهة ويب مركزية."
    ]
    for m in mech_1:
        add_paragraph(doc, m)
    
    add_paragraph(doc, "رابط المشروع (للاطلاع والصور): https://github.com/mydlp/mydlp", bold=True)
    add_paragraph(doc, "[صورة 1: واجهة لوحة التحكم الرئيسية في نظام MyDLP]")
    doc.add_paragraph()

    # Study 2: PacketFence
    add_heading(doc, "2.2.2 مشروع (PacketFence): التحكم في الوصول وأمن الشبكة", level=3)
    add_paragraph(doc, "هو حل شامل للتحكم في الوصول للشبكة (NAC)، ولكنه يتضمن ميزات قوية لمراقبة السلوك واكتشاف التهديدات الداخلية، مما يجعله نموذجاً ممتازاً لكيفية حماية البنية التحتية للبيانات.")
    
    add_paragraph(doc, "آلية عمل النظام:", bold=True)
    mech_2 = [
        "1. التسجيل والمراقبة: يفرض النظام على كل جهاز التسجيل قبل الدخول للشبكة، ويراقب نشاطه باستمرار.",
        "2. كشف التسلل (IDS): يستخدم أدوات مثل Snort/Suricata لتحليل حركة الشبكة واكتشاف الأنماط المشبوهة التي قد تشير إلى محاولة تسريب بيانات.",
        "3. العزل التلقائي (Isolation): عند اكتشاف تهديد، يقوم النظام تلقائياً بعزل الجهاز المصاب في شبكة افتراضية (VLAN) معزولة لمنع انتشار الضرر أو تسرب البيانات.",
        "4. الاستجابة: يرسل تنبيهات للمسؤولين مع تفاصيل الحادثة لاتخاذ الإجراءات التصحيحية."
    ]
    for m in mech_2:
        add_paragraph(doc, m)

    add_paragraph(doc, "رابط المشروع (للاطلاع والصور): https://www.packetfence.org/about.html", bold=True)
    add_paragraph(doc, "[صورة 2: واجهة إدارة الأجهزة والمخاطر في PacketFence]")
    doc.add_paragraph()

    # Study 3: Wazuh
    add_heading(doc, "2.2.3 مشروع (Wazuh): منصة الأمن السيبراني الموحدة", level=3)
    add_paragraph(doc, "يُعد Wazuh منصة أمنية مفتوحة المصدر (Open Source Security Platform) توفر حماية شاملة للنقاط النهائية (Endpoints) والبنية التحتية السحابية. يجمع بين قدرات XDR و SIEM، ويتميز بقدرته العالية على مراقبة سلامة الملفات (FIM) والامتثال للسياسات.")
    
    add_paragraph(doc, "آلية عمل النظام:", bold=True)
    mech_3 = [
        "1. الوكيل (Agent): يتم تثبيت وكيل خفيف على الأجهزة (Windows/Linux) لجمع السجلات ومراقبة الملفات وتنفيذ الأوامر.",
        "2. الخادم (Manager): يقوم باستقبال البيانات من الوكلاء، وتحليلها باستخدام قواعد (Decoders & Rules) لاكتشاف التهديدات.",
        "3. الفهرسة (Indexer): يتم فهرسة البيانات وتخزينها للبحث السريع.",
        "4. الواجهة (Dashboard): واجهة رسومية متقدمة (مبنية على Kibana/OpenSearch) لعرض التنبيهات وإدارة الحوادث الأمنية."
    ]
    for m in mech_3:
        add_paragraph(doc, m)

    add_paragraph(doc, "رابط المشروع (للصور): https://wazuh.com/platform/", bold=True)
    add_paragraph(doc, "[صورة 3: لوحة تحكم Wazuh الرئيسية]")
    doc.add_paragraph()
    
    # Comparison Table (Features/Weaknesses)
    add_paragraph(doc, "جدول (2-1): تحليل الدراسات السابقة (المميزات والعيوب)", bold=True)
    lit_table = doc.add_table(rows=1, cols=3)
    lit_table.style = 'Table Grid'
    
    l_hdr = lit_table.rows[0].cells
    l_hdr[0].text = 'النظام (المشروع)'
    l_hdr[1].text = 'المميزات (Strengths)'
    l_hdr[2].text = 'العيوب/الفجوة (Weaknesses)'
    
    for cell in l_hdr:
        for paragraph in cell.paragraphs:
            set_rtl(paragraph)
            for run in paragraph.runs:
                run.font.bold = True
                set_run_rtl(run)

    studies_data = [
        (
            "MyDLP",
            "قوي جداً في مراقبة الشبكة (Network) والمنع الفوري للبروتوكولات (Web/Mail).",
            "يعتمد بشكل كبير على القواعد الثابتة (Regex) ويفتقر للدقة السياقية التي يوفرها الذكاء الاصطناعي."
        ),
        (
            "PacketFence",
            "تحكم ممتاز في الوصول وعزل الأجهزة المصابة تلقائياً.",
            "يركز على أمن الشبكة (Network Security) أكثر من تركيزه على فحص محتوى البيانات (Content Inspection)."
        ),
        (
            "Wazuh",
            "منصة شاملة وقوية جداً للمراقبة والامتثال (Compliance) مع واجهة بصرية مذهلة.",
            "يركز على اكتشاف التهديدات (Threat Detection) وسلامة الملفات، وليس متخصصاً في منع تسرب البيانات (DLP) بمعناه الدقيق."
        )
    ]

    for item in studies_data:
        row_cells = lit_table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                set_rtl(paragraph)
                for run in paragraph.runs:
                    set_run_rtl(run)
    doc.add_paragraph()

    # 2.3 Proposal System
    add_heading(doc, "2.3 النظام المقترح (Proposal System)", level=2)
    proposal_intro = """يقدم هذا المشروع نظام "Athier"، وهو منصة أمنية متكاملة تهدف إلى سد الفجوة التقنية في أنظمة حماية البيانات الحالية. يعتمد النظام المقترح على نهج استباقي (Proactive Approach) يجمع بين دقة الذكاء الاصطناعي في فهم النصوص وبين صرامة قواعد الحماية التقليدية.
    
فكرة النظام الأساسية هي عدم الاكتفاء بمطابقة الأنماط الشكلية للبيانات (مثل عدد الخانات في رقم البطاقة)، بل تحليل "السياق الدلالي" للنص باستخدام نماذج معالجة اللغات الطبيعية (NLP). هذا يسمح للنظام بالتمييز بين البيانات الحساسة الحقيقية والبيانات المشابهة لها، مما يرفع كفاءة العمليات الأمنية ويقلل من تعطيل سير العمل الناتج عن الإنذارات الخاطئة."""
    for paragraph in proposal_intro.split('\n\n'):
        add_paragraph(doc, paragraph.strip())
    
    # Comparison Table
    add_paragraph(doc, "جدول (2-2): مقارنة شاملة بين الأنظمة السابقة والنظام المقترح", bold=True)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'وجه المقارنة'
    hdr_cells[1].text = 'MyDLP'
    hdr_cells[2].text = 'PacketFence'
    hdr_cells[3].text = 'Wazuh'
    hdr_cells[4].text = 'النظام المقترح (Athier)'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            set_rtl(paragraph)
            for run in paragraph.runs:
                run.font.bold = True
                set_run_rtl(run)

    comparison_data = [
        ("التركيز الأساسي", "مراقبة الشبكة (Network)", "التحكم في الوصول (NAC)", "أمن النقاط النهائية (Endpoint/SIEM)", "حماية شاملة ذكية (Context-aware DLP)"),
        ("آلية الاكتشاف", "قواعد ثابتة (Regex/Fingerprint)", "سلوك الشبكة (Behavior)", "تحليل السجلات (Log Analysis)", "الذكاء الاصطناعي (AI/NLP) + Regex"),
        ("دقة الاكتشاف", "متوسطة (يعتمد على القواعد)", "عالية في الشبكة / منخفضة في المحتوى", "عالية في التهديدات الأمنية", "عالية جداً (يفهم السياق واللغة)"),
        ("الإجراءات الوقائية", "منع، حجر، تنبيه", "عزل الجهاز (VLAN Isolation)", "تنبيه، استجابة نشطة (Active Response)", "منع، تشفير، تنبيه، عزل (متعدد)"),
        ("دعم اللغة العربية", "ضعيف / غير مدعوم", "غير مدعوم (واجهة فقط)", "يدعم الواجهة (Unicode)", "دعم كامل (نماذج NLP عربية)"),
        ("سهولة الاستخدام", "واجهة ويب تقليدية", "معقدة (للمحترفين)", "واجهة متقدمة لكنها معقدة", "واجهة عصرية تفاعلية (Dashboard)")
    ]

    for item in comparison_data:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]
        row_cells[4].text = item[4]
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                set_rtl(paragraph)
                for run in paragraph.runs:
                    set_run_rtl(run)
    doc.add_paragraph()

    # 2.4 System Overview
    add_heading(doc, "2.4 النظرة العامة للنظام (System Overview)", level=2)
    overview_text = """يتكون نظام "Athier" من بنية خدمية (Microservices Architecture) مرنة تضمن القابلية للتوسع والصيانة. يتألف النظام من المكونات الرئيسية التالية:

1. واجهة المستخدم (Frontend): وهي بوابة التفاعل مع المسؤولين، تم بناؤها باستخدام تقنيات الويب الحديثة لتوفير لوحة تحكم تفاعلية تعرض التنبيهات، الإحصائيات، وتسمح بإدارة السياسات الأمنية بسهولة.

2. وحدة إدارة الـ API (Backend API): تمثل العقل المدبر للنظام، حيث تستقبل الطلبات من الواجهة أو من نقاط المراقبة، وتقوم بتوجيهها إلى الخدمات المختصة. تم تطويرها باستخدام إطار عمل FastAPI لضمان سرعة الاستجابة.

3. خدمة التحليل الذكي (Analysis Service): تعتمد على محرك Microsoft Presidio، وهي المسؤولة عن استقبال النصوص وفحصها باستخدام نماذج تعلم الآلة المدربة مسبقاً (Pre-trained Models) لاكتشاف الكيانات الحساسة مثل الأسماء، العناوين، وأرقام الهويات.

4. محرك السياسات (Policy Engine): يقوم بمقارنة النتائج الواردة من خدمة التحليل مع القواعد التي وضعها المسؤول (مثل: "اكتشاف أرقام بطاقات ائتمان" + "مستوى خطورة مرتفع" = "منع").

5. قاعدة البيانات (Database): مستودع مركزي آمن (PostgreSQL) لتخزين سجلات التدقيق (Logs)، تفاصيل التنبيهات، وتكوينات النظام."""
    for paragraph in overview_text.split('\n\n'):
        add_paragraph(doc, paragraph.strip())
    doc.add_paragraph()

    # 2.5 System Working Procedure
    add_heading(doc, "2.5 آلية عمل النظام (System Working Procedure)", level=2)
    add_paragraph(doc, "تتم عملية حماية البيانات داخل النظام عبر سلسلة من الخطوات المتتابعة والآلية، والتي تضمن معالجة البيانات في الوقت الفعلي:")
    
    steps = [
        "1. اعتراض البيانات (Interception): يقوم النظام (عبر نقاط المراقبة) باعتراض البيانات المرسلة (مثل محتوى بريد إلكتروني أو ملف) قبل خروجها من الشبكة الداخلية.",
        "2. الاستخراج والمعالجة (Extraction): يتم استخراج النص الصرف من البيانات المعترضة وتجهيزه للتحليل (تنظيف النص وإزالة التنسيقات).",
        "3. التحليل الذكي (Intelligent Analysis): يُرسل النص إلى خدمة Presidio التي تقوم بمسحه بحثاً عن كيانات محددة (PII). يعيد المحرك قائمة بالنتائج مع 'درجة الثقة' (Confidence Score) لكل نتيجة.",
        "4. مطابقة السياسات (Policy Matching): يتحقق النظام مما إذا كانت البيانات المكتشفة تخالف أياً من السياسات النشطة. مثلاً، إذا تم اكتشاف 'رقم هاتف' وكانت هناك سياسة تمنع مشاركة أرقام الهواتف.",
        "5. اتخاذ الإجراء (Enforcement): بناءً على المطابقة، ينفذ النظام الإجراء المحدد (Block لمنع الإرسال، Alert لإرسال تنبيه فقط، أو Encrypt لتشفير البيانات).",
        "6. التوثيق (Auditing): يتم تسجيل تفاصيل العملية كاملة (الوقت، المستخدم، نوع البيانات، الإجراء المتخذ) في قاعدة البيانات لأغراض التدقيق والمراجعة اللاحقة."
    ]
    for step in steps:
        add_paragraph(doc, step)
    doc.add_paragraph()

    # 2.6 Feasibility Study
    add_heading(doc, "2.6 دراسة الجدوى (Feasibility Study)", level=2)
    add_paragraph(doc, "تم إجراء دراسة جدوى شاملة للتأكد من إمكانية تنفيذ المشروع ونجاحه من النواحي التقنية والتشغيلية والاقتصادية:")
    
    # 2.6.1 Technical Feasibility
    add_paragraph(doc, "2.6.1 دراسة الجدوى التقنية (Technical Feasibility):", bold=True)
    add_paragraph(doc, "تعتبر الجدوى التقنية الركيزة الأساسية للمشروع، حيث تم التأكد من توفر جميع الأدوات والتقنيات اللازمة لبناء النظام. يعتمد النظام على بيئة تطوير حديثة ومفتوحة المصدر.")
    
    add_paragraph(doc, "جدول (2-4): المتطلبات التقنية والمعدات واستخداماتها", bold=True)
    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.style = 'Table Grid'
    t_hdr = tech_table.rows[0].cells
    t_hdr[0].text = 'التقنية / المعدات'
    t_hdr[1].text = 'الوصف / النوع'
    t_hdr[2].text = 'الاستخدام في النظام'
    
    for cell in t_hdr:
        for paragraph in cell.paragraphs:
            set_rtl(paragraph)
            for run in paragraph.runs:
                run.font.bold = True
                set_run_rtl(run)

    tech_data = [
        ("لغة البرمجة", "Python 3.11+", "تطوير الواجهة الخلفية (Backend) وسكربتات المعالجة."),
        ("إطار العمل", "FastAPI", "بناء الـ API السريع للربط بين الواجهة والخدمات."),
        ("محرك التحليل", "Microsoft Presidio", "تحليل النصوص واكتشاف البيانات الحساسة (PII)."),
        ("قاعدة البيانات", "PostgreSQL", "تخزين السياسات، السجلات، وتفاصيل التنبيهات."),
        ("بيئة التشغيل", "Docker & Docker Compose", "تغليف النظام لضمان عمله على أي خادم."),
        ("جهاز الخادم (Server)", "Core i5 / 16GB RAM", "استضافة النظام ومعالجة طلبات التحليل."),
        ("نظام التشغيل", "Linux (Ubuntu) / Windows", "البيئة الأساسية لتشغيل الخادم.")
    ]

    for item in tech_data:
        row_cells = tech_table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                set_rtl(paragraph)
                for run in paragraph.runs:
                    set_run_rtl(run)
    doc.add_paragraph()

    # 2.6.2 Operational Feasibility
    add_paragraph(doc, "2.6.2 دراسة الجدوى التشغيلية (Operational Feasibility):", bold=True)
    add_paragraph(doc, "يركز هذا الجانب على كيفية تشغيل النظام والاستفادة منه في البيئة الفعلية. تم تصميم النظام ليكون:")
    op_points = [
        "سهل الاستخدام: واجهة ويب بسيطة لا تتطلب تدريباً تقنياً معقداً لموظفي الأمن.",
        "مؤتمت بالكامل: يعمل النظام في الخلفية (Background Service) ويقوم بالتحليل واتخاذ القرار آلياً دون تدخل بشري، مما يرفع الكفاءة التشغيلية.",
        "قابل للصيانة: استخدام Docker يجعل عملية التحديث والصيانة سهلة وسريعة دون توقف طويل للخدمة.",
        "تحسين سير العمل: بدلاً من المراجعة اليدوية البطيئة للملفات، يقوم النظام بفلترة البيانات وتنبيه المسؤول فقط في الحالات الحرجة."
    ]
    for p in op_points:
        add_bullet_point(doc, p)
    doc.add_paragraph()

    # 2.6.3 Economic Feasibility
    add_paragraph(doc, "2.6.3 دراسة الجدوى الاقتصادية (Economic Feasibility):", bold=True)
    add_paragraph(doc, "توضح الدراسة الاقتصادية التكاليف التقريبية للمشروع، والتي تعتبر منخفضة جداً مقارنة بالأنظمة التجارية، نظراً للاعتماد على المصادر المفتوحة.")

    add_paragraph(doc, "جدول (2-5): التكاليف التقديرية للمشروع", bold=True)
    eco_table = doc.add_table(rows=1, cols=3)
    eco_table.style = 'Table Grid'
    e_hdr = eco_table.rows[0].cells
    e_hdr[0].text = 'البند'
    e_hdr[1].text = 'التكلفة التقديرية'
    e_hdr[2].text = 'ملاحظات'
    
    for cell in e_hdr:
        for paragraph in cell.paragraphs:
            set_rtl(paragraph)
            for run in paragraph.runs:
                run.font.bold = True
                set_run_rtl(run)

    eco_data = [
        ("تراخيص البرمجيات (Software Licenses)", "0 $", "جميع الأدوات (Python, Presidio, Postgres) مفتوحة المصدر."),
        ("أجهزة الخادم (Server Hardware)", "500 $ - 800 $", "تكلفة جهاز كمبيوتر متوسط المواصفات (أو استخدام خادم موجود)."),
        ("تطوير وبرمجة (Development)", "جهد ذاتي", "قام فريق العمل بالتطوير (مشروع تخرج)."),
        ("الصيانة والتشغيل (سنوياً)", "50 $", "تكاليف كهرباء / انترنت تقديرية."),
        ("الإجمالي", "~ 550 $", "تكلفة تأسيسية منخفضة جداً.")
    ]

    for item in eco_data:
        row_cells = eco_table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                set_rtl(paragraph)
                for run in paragraph.runs:
                    set_run_rtl(run)
    doc.add_paragraph()

    # 2.7 Risks Management
    add_heading(doc, "2.7 إدارة المخاطر (Risks Management)", level=2)
    
    # 2.7.1 Risk Identification
    add_paragraph(doc, "2.7.1 تحديد المخاطر (Risk Identification):", bold=True)
    add_paragraph(doc, "تم تحليل المشروع لتحديد المخاطر المحتملة وتصنيفها حسب نوعها ودرجة تأثيرها، كما هو موضح في سجل المخاطر التالي:")
    
    # Risk Register Table
    add_paragraph(doc, "جدول (2-6): سجل تحديد وتصنيف المخاطر", bold=True)
    reg_table = doc.add_table(rows=1, cols=5)
    reg_table.style = 'Table Grid'
    
    reg_hdr = reg_table.rows[0].cells
    reg_hdr[0].text = 'وصف الخطر'
    reg_hdr[1].text = 'نوع الخطر'
    reg_hdr[2].text = 'الاحتمالية'
    reg_hdr[3].text = 'الأثر (Impact)'
    reg_hdr[4].text = 'مستوى الخطر' # (High/Medium/Low)
    
    for cell in reg_hdr:
        for paragraph in cell.paragraphs:
            set_rtl(paragraph)
            for run in paragraph.runs:
                run.font.bold = True
                set_run_rtl(run)

    register_data = [
        ("بطء النظام عند تحليل نصوص ضخمة", "تقني (Technical)", "متوسطة", "عالي", "مرتفع (High)"),
        ("عدم دقة النموذج في اكتشاف اللهجات", "تقني (Technical)", "عالية", "متوسط", "مرتفع (High)"),
        ("توقف الخادم المفاجئ", "تشغيلي (Operational)", "منخفضة", "حرج", "متوسط (Medium)"),
        ("صعوبة ربط النظام بالشبكة القديمة", "تقني (Integration)", "متوسطة", "متوسط", "متوسط (Medium)"),
        ("تعرض قاعدة البيانات للاختراق", "أمني (Security)", "منخفضة", "حرج", "مرتفع (High)")
    ]

    for item in register_data:
        row_cells = reg_table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]
        row_cells[4].text = item[4]
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                set_rtl(paragraph)
                for run in paragraph.runs:
                    set_run_rtl(run)
    doc.add_paragraph()

    # 2.7.2 Risk Management Plan
    add_paragraph(doc, "2.7.2 خطة مواجهة المخاطر (Risk Management Plan):", bold=True)
    add_paragraph(doc, "بناءً على سجل المخاطر أعلاه، تم وضع استراتيجيات للتعامل مع المخاطر ذات المستوى المرتفع والمتوسط:")

    # Risk Mitigation Table (Renumbered to 2-7)
    add_paragraph(doc, "جدول (2-7): خطة مواجهة المخاطر واستراتيجية المعالجة", bold=True)
    risk_table = doc.add_table(rows=1, cols=3) # Simplified cols since Impact/Prob is in prev table
    risk_table.style = 'Table Grid'
    
    r_hdr = risk_table.rows[0].cells
    r_hdr[0].text = 'الخطر'
    r_hdr[1].text = 'استراتيجية المعالجة (Mitigation Strategy)'
    r_hdr[2].text = 'الإجراء البديل (Contingency Plan)'
    
    for cell in r_hdr:
        for paragraph in cell.paragraphs:
            set_rtl(paragraph)
            for run in paragraph.runs:
                run.font.bold = True
                set_run_rtl(run)

    risk_data = [
        ("بطء النظام", "استخدام المعالجة غير المتزامنة (Async)", "زيادة موارد الخادم (Vertical Scaling)"),
        ("عدم دقة النموذج", "تدريب النموذج على بيانات محلية إضافية", "تفعيل قواعد Regex يدوية كبديل مؤقت"),
        ("توقف الخادم", "استخدام Docker و Kubernetes لضمان التوافر", "تشغيل خادم احتياطي (Standby Server)"),
        ("صعوبة الربط", "بناء API قياسي (RESTful)", "استخدام وسطاء (Middlewares) للتحويل"),
        ("اختراق البيانات", "تشفير البيانات الحساسة (AES-256)", "استعادة النسخة الاحتياطية فوراً")
    ]

    for item in risk_data:
        row_cells = risk_table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                set_rtl(paragraph)
                for run in paragraph.runs:
                    set_run_rtl(run)

    # Save document
    output_path = 'Graduation_Project_Chapter2_v14.docx'
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    create_document()
