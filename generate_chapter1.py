import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_rtl(paragraph):
    """Sets the paragraph direction to RTL."""
    pPr = paragraph._p.get_or_add_pPr()
    section = OxmlElement('w:bidi')
    section.set(qn('w:val'), '1')
    pPr.append(section)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_run_rtl(run):
    """Sets the run direction to RTL (for mixed text)."""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    set_rtl(heading)
    for run in heading.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(16 if level == 1 else 14)
        run.font.color.rgb = RGBColor(0, 0, 0) # Black
        set_run_rtl(run)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    set_run_rtl(run)
    return p

def add_bullet_point(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_run_rtl(run)

def create_document():
    doc = Document()

    # Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # --- Preliminary Pages ---
    # 1. Dedication
    doc.add_page_break()
    add_heading(doc, "I. الإهداء (Dedication)", level=1)
    dedication_text = """إلى من أضاءوا لي الطريق...
إلى والديّ العزيزين، مصدر قوتي وإلهامي.
إلى أساتذتي الأفاضل الذين لم يبخلوا عليّ بعلمهم.
إلى زملائي وأصدقائي الذين شاركوني هذه الرحلة.
أهدي هذا العمل المتواضع، راجياً أن يكون خطوة نافعة في مسيرة العلم."""
    for line in dedication_text.split('\n'):
        add_paragraph(doc, line)

    # 2. Authorization
    doc.add_page_break()
    add_heading(doc, "II. التفويض (Authorization)", level=1)
    auth_text = """أفوض أنا الطالب/الطالبة (الاسم) جامعة الرازي باستخدام هذا المشروع للأغراض الأكاديمية والبحثية، مع الاحتفاظ بحقوق الملكية الفكرية للأفكار الأصيلة الواردة فيه.

التوقيع: ____________________
التاريخ: ____________________"""
    for line in auth_text.split('\n'):
        add_paragraph(doc, line)

    # 3. Acknowledgment
    doc.add_page_break()
    add_heading(doc, "III. شكر وتقدير (Acknowledgment)", level=1)
    ack_text = """الحمد لله رب العالمين، والصلاة والسلام على أشرف الأنبياء والمرسلين.
أتقدم بجزيل الشكر وعظيم الامتنان إلى المشرف على هذا المشروع، الدكتور الفاضل (اسم المشرف)، على توجيهاته السديدة ودعمه المستمر.
كما أشكر كلية الحاسوب وتكنولوجيا المعلومات بجامعة الرازي على توفير البيئة الأكاديمية المحفزة.
والشكر موصول لكل من ساهم، ولو بكلمة طيبة، في إنجاز هذا العمل."""
    for line in ack_text.split('\n'):
        add_paragraph(doc, line)

    # 4. Examiner Committee
    doc.add_page_break()
    add_heading(doc, "IV. لجنة المناقشة (Examiner Committee)", level=1)
    exam_text = """تمت مناقشة هذا المشروع وإجازته من قبل اللجنة المكونة من:

1. الدكتور/ة: ____________________ (رئيساً ومشرفاً)
2. الدكتور/ة: ____________________ (مناقشاً داخلياً)
3. الدكتور/ة: ____________________ (مناقشاً خارجياً)

تاريخ المناقشة: __/__/2024"""
    for line in exam_text.split('\n'):
        add_paragraph(doc, line)

    # 5. Abstract
    doc.add_page_break()
    add_heading(doc, "V. الملخص (Abstract)", level=1)
    abstract_text = """يهدف هذا المشروع إلى بناء نظام متكامل لحماية البيانات (Integrated Data Protection System) يجمع بين تقنيات منع فقدان البيانات (DLP) وتحليل النصوص باستخدام الذكاء الاصطناعي. يعالج المشروع مشكلة تسريب البيانات الحساسة (مثل أرقام الهواتف، والبيانات الشخصية PII) داخل المؤسسات، حيث يعتمد على أدوات مفتوحة المصدر قوية مثل Microsoft Presidio للتحليل الدلالي و MyDLP للمراقبة. يوفر النظام واجهة ويب تفاعلية لإدارة السياسات الأمنية، ومراقبة حركة البيانات في الوقت الفعلي، وإصدار تنبيهات فورية عند اكتشاف أي انتهاك. تم تطوير النظام باستخدام لغة Python وإطار عمل FastAPI، مع قاعدة بيانات PostgreSQL لتخزين السجلات والسياسات، مما يجعله حلاً قابلاً للتوسع وفعالاً من حيث التكلفة للمؤسسات الصغيرة والمتوسطة."""
    add_paragraph(doc, abstract_text)

    # 6. Table of Contents
    doc.add_page_break()
    add_heading(doc, "VI. فهرس المحتويات (Table of Contents)", level=1)
    add_paragraph(doc, "[يتم توليده تلقائياً في Word]")

    # 7. List of Figures
    doc.add_page_break()
    add_heading(doc, "VII. قائمة الأشكال (List of Figures)", level=1)
    add_paragraph(doc, "[قائمة الأشكال]")

    # 8. List of Tables
    doc.add_page_break()
    add_heading(doc, "VIII. قائمة الجداول (List of Tables)", level=1)
    add_paragraph(doc, "[قائمة الجداول]")

    # 9. List of Abbreviations
    doc.add_page_break()
    add_heading(doc, "IX. قائمة الاختصارات (List of Abbreviations)", level=1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'الاختصار (Abbreviation)'
    hdr_cells[1].text = 'المعنى (Meaning)'
    
    # Set header to RTL
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            set_rtl(paragraph)
            for run in paragraph.runs:
                run.font.bold = True
                set_run_rtl(run)

    abbreviations = [
        ("DLP", "Data Loss Prevention (منع فقدان البيانات)"),
        ("PII", "Personally Identifiable Information (معلومات تحديد الهوية الشخصية)"),
        ("GDPR", "General Data Protection Regulation (اللائحة العامة لحماية البيانات)"),
        ("HIPAA", "Health Insurance Portability and Accountability Act"),
        ("API", "Application Programming Interface"),
        ("NLP", "Natural Language Processing (معالجة اللغات الطبيعية)"),
        ("UI", "User Interface (واجهة المستخدم)")
    ]

    for abbr, meaning in abbreviations:
        row_cells = table.add_row().cells
        row_cells[0].text = abbr
        row_cells[1].text = meaning
        
        # Set row cells to RTL
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                set_rtl(paragraph)
                for run in paragraph.runs:
                    set_run_rtl(run)

    doc.add_page_break()

    # --- Chapter One ---
    add_heading(doc, "ثانيًا: الفصل الأول – (Chapter One)", level=1)
    doc.add_paragraph()

    # 1. Introduction
    add_heading(doc, "1. مقدمة (Introduction)", level=2)
    text_intro = """يشهد العالم الرقمي اليوم تزايداً هائلاً في حجم البيانات المتداولة عبر الشبكات، مما جعل أمن المعلومات وحماية البيانات الشخصية (Data Privacy) أحد أهم التحديات التي تواجه المؤسسات والشركات. في ظل تصاعد الهجمات السيبرانية وحوادث تسريب البيانات، لم تعد حماية البيانات مجرد خيار تقني، بل أصبحت ضرورة استراتيجية وقانونية ملحة لضمان استمرارية الأعمال والحفاظ على سمعة المؤسسات.

تُعاني العديد من الأنظمة التقليدية من قصور في التمييز الدقيق بين البيانات العادية والبيانات الحساسة، مما يؤدي إما إلى تسرب بيانات حرجة أو إلى تعطيل العمل بسبب الإنذارات الخاطئة. من هنا برزت الحاجة إلى أنظمة متطورة لمنع فقدان البيانات (DLP - Data Loss Prevention) تعتمد على تقنيات الذكاء الاصطناعي ومعالجة اللغات الطبيعية (NLP) لتقديم حماية دقيقة وفعالة."""
    for paragraph in text_intro.split('\n\n'):
        add_paragraph(doc, paragraph.strip())
    doc.add_paragraph()

    # 2. Definition
    add_heading(doc, "2. تعريف المشروع (Definition of the Project)", level=2)
    text_def = """مشروع "Secure" هو نظام متكامل لحماية البيانات (Integrated Data Protection System) مصمم لاكتشاف ومنع تسرب البيانات الحساسة والمعلومات الشخصية (PII) داخل المؤسسات.

النظام عبارة عن تطبيق ويب (Web Application) متكامل يجمع بين قوة الذكاء الاصطناعي في تحليل النصوص وآليات الحماية التقليدية، ويستهدف مسؤولي أمن المعلومات في المؤسسات المتوسطة والكبيرة التي تتعامل مع بيانات حساسة للعملاء أو الموظفين. يوفر النظام واجهة تحكم مركزية لإدارة سياسات الحماية، مراقبة حركة البيانات، واتخاذ إجراءات تلقائية لمنع التسريب."""
    for paragraph in text_def.split('\n\n'):
        add_paragraph(doc, paragraph.strip())
    doc.add_paragraph()

    # 3. Problem Statement
    add_heading(doc, "3. تحديد المشكلة (Problem Statement)", level=2)
    add_paragraph(doc, "تواجه المؤسسات حالياً فجوة كبيرة بين أدوات الكشف عن البيانات وأدوات منع التسريب، وتتلخص المشكلة في النقاط التالية:")
    
    problems = [
        "قصور أنظمة DLP التقليدية: تعتمد معظم الأنظمة الحالية على التعبيرات النمطية (Regex) البسيطة، مما يؤدي إلى نسبة عالية من الإنذارات الكاذبة (False Positives)، حيث قد يتم حظر مستندات عادية لمجرد احتواءها على أرقام تشبه أرقام البطاقات الائتمانية.",
        "ضعف الامتثال للمعايير الدولية: تجد المؤسسات صعوبة في إثبات الامتثال لقوانين حماية البيانات الصارمة مثل (GDPR) و(HIPAA) بسبب غياب سجلات دقيقة ومفصلة توضح نوع البيانات التي تم حمايتها وكيفية التعامل معها.",
        "تشتت أدوات الحماية: تضطر المؤسسات غالباً لاستخدام عدة أدوات منفصلة (أداة للمسح، وأخرى للمنع، وثالثة للتقارير)، مما يزيد من تعقيد الإدارة ويرفع التكلفة التشغيلية ويخلق ثغرات أمنية محتملة."
    ]
    for p in problems:
        add_bullet_point(doc, p)
        
    add_paragraph(doc, "يعالج هذا المشروع هذه الثغرات من خلال دمج محرك تحليل ذكي (Microsoft Presidio) مع نظام منع تسريب (MyDLP) في منصة موحدة توفر دقة عالية في الاكتشاف وسهولة في الإدارة.")
    doc.add_paragraph()

    # 4. Objectives
    add_heading(doc, "4. أهداف المشروع (Project Objectives)", level=2)
    
    add_paragraph(doc, "الهدف العام (Main Objective):", bold=True)
    add_paragraph(doc, "تطوير نظام أمني متكامل وذكي لحماية البيانات الحساسة (PII) ومنع تسربها باستخدام تقنيات معالجة اللغات الطبيعية، بما يضمن الامتثال لمعايير الخصوصية العالمية.")
    
    add_paragraph(doc, "الأهداف الفرعية (Specific Objectives):", bold=True)
    objectives = [
        "دمج مكتبة Microsoft Presidio لتحليل النصوص واكتشاف الكيانات الحساسة (مثل الأسماء، أرقام الهواتف، البطاقات الائتمانية) بدقة عالية وتقليل الإنذارات الكاذبة.",
        "تطوير وحدة إدارة سياسات مرنة تسمح للمسؤولين بتحديد إجراءات الحماية (منع، تنبيه، تشفير) بناءً على نوع البيانات ومستوى خطورتها.",
        "بناء لوحة تحكم (Dashboard) تفاعلية لعرض إحصائيات التهديدات وحالة النظام في الوقت الفعلي.",
        "توفير آلية لتشفير البيانات الحساسة المكتشفة (Encryption) قبل تخزينها لضمان أقصى درجات الأمان.",
        "إنشاء نظام سجلات وتقارير شامل (Auditing & Reporting) يدعم متطلبات الامتثال لقوانين GDPR و HIPAA."
    ]
    for obj in objectives:
        add_bullet_point(doc, obj)
    doc.add_paragraph()

    # 5. Importance
    add_heading(doc, "5. أهمية المشروع (Project Important)", level=2)
    add_paragraph(doc, "تكمن أهمية المشروع في الجوانب التالية:")
    importance = [
        "للمؤسسات: حماية الأصول المعرفية والبيانات السرية من التسريب، مما يحمي سمعة المؤسسة ويجنبها الخسائر المالية والقانونية.",
        "عملياً وتقنياً: تقديم نموذج تطبيقي لكيفية دمج تقنيات الذكاء الاصطناعي (NLP) في أنظمة الأمن السيبراني التقليدية لرفع كفاءتها، مما يقلل الجهد اليدوي المطلوب من فرق الأمن لمراجعة التنبيهات.",
        "قانونياً: مساعدة المؤسسات على تجنب الغرامات الباهظة المترتبة على مخالفة لوائح حماية البيانات (مثل لائحة حماية البيانات العامة GDPR)."
    ]
    for item in importance:
        add_bullet_point(doc, item)
    doc.add_paragraph()

    # 6. Scope and Limitations
    add_heading(doc, "6. نطاق وحدود المشروع (Project Scope and Limitations)", level=2)
    
    add_paragraph(doc, "نطاق المشروع (Scope):", bold=True)
    scope = [
        "يغطي المشروع تحليل النصوص (Text Analysis) للكشف عن أنماط البيانات الشخصية (PII).",
        "إدارة ومراقبة حركة البيانات (Data Traffic Monitoring) عبر محاكاة لبيئة شبكية أو تكامل مع خوادم البريد الإلكتروني.",
        "يدعم النظام اللغة الإنجليزية بشكل أساسي مع قابلية لدعم لغات أخرى.",
        "إدارة المستخدمين والصلاحيات والسياسات الأمنية."
    ]
    for item in scope:
        add_bullet_point(doc, item)
        
    add_paragraph(doc, "حدود المشروع (Limitations):", bold=True)
    limitations = [
        "البيئة: تم تطوير النظام واختباره في بيئة معملية، وقد يحتاج إلى تعديلات في البنية التحتية عند نشره في بيئات إنتاج ضخمة (Large-scale Enterprise).",
        "أنواع الملفات: يركز النظام حالياً على البيانات النصية، ولا يشمل تحليل الصور (OCR) أو ملفات الفيديو المشفرة في الإصدار الحالي.",
        "الأداء: يعتمد زمن المعالجة على قدرات الخادم، وقد يتأثر عند تحليل كميات ضخمة جداً من البيانات في الوقت الفعلي."
    ]
    for item in limitations:
        add_bullet_point(doc, item)
    doc.add_paragraph()

    # 7. Tools
    add_heading(doc, "7. الأدوات المستخدمة في المشروع (Tools used in the Project)", level=2)
    
    tools = {
        "لغات البرمجة": [
            "Python: تم اختيارها لمرونتها العالية، وتوفر مكتبات قوية في مجال الأمن السيبراني والذكاء الاصطناعي، وهي اللغة الأساسية للواجهة الخلفية (Backend).",
            "JavaScript (HTML/CSS): لبناء واجهة المستخدم التفاعلية (Frontend)."
        ],
        "أطر العمل (Frameworks)": [
            "FastAPI: لبناء واجهة برمجية (API) سريعة وعالية الأداء، وتوفير توثيق تلقائي للنظام.",
            "Microsoft Presidio: إطار عمل متخصص من مايكروسوفت لاكتشاف وإخفاء المعلومات الحساسة، تم استخدامه لدقته وموثوقيته."
        ],
        "قواعد البيانات": [
            "PostgreSQL: قاعدة بيانات علائقية قوية ومفتوحة المصدر، استُخدمت لتخزين السجلات، السياسات، وبيانات المستخدمين نظراً لموثوقيتها في البيئات الإنتاجية.",
            "SQLAlchemy (ORM): لتسهيل التعامل مع قاعدة البيانات بطريقة برمجية آمنة."
        ],
        "أدوات أخرى": [
            "Docker & Docker Compose: لضمان سهولة نشر النظام وتوحيد بيئة التشغيل (Containerization).",
            "Cryptography Library: لتنفيذ خوارزميات التشفير (AES) لحماية البيانات المخزنة.",
            "MyDLP (Integration): تم دمج مفاهيم ووظائف MyDLP لتنفيذ سياسات المنع والمراقبة."
        ]
    }
    
    for category, items in tools.items():
        add_paragraph(doc, category + ":", bold=True)
        for item in items:
            add_bullet_point(doc, item)
    doc.add_paragraph()

    # 8. Methodology
    add_heading(doc, "8. المنهجية المستخدمة لعمل المشروع (Research Methodology)", level=2)
    add_paragraph(doc, "تم اتباع منهجية دورة حياة تطوير النظم (SDLC) بأسلوب تكراري (Iterative Approach) لضمان الجودة والمرونة، وتضمنت المراحل التالية:")
    
    methodology_steps = [
        "جمع المتطلبات (Requirements Analysis): دراسة احتياجات أنظمة الـ DLP وتحليل الثغرات في الحلول الحالية (Presidio, MyDLP).",
        "التصميم (System Design): تصميم هيكلية النظام (Architecture)، قاعدة البيانات (ERD)، وتدفق البيانات (Data Flow).",
        "التنفيذ (Implementation): برمجة الخدمات الأساسية (خدمة التحليل، خدمة السياسات، خدمة التشفير) وربطها ببعضها البعض.",
        "الاختبار (Testing): إجراء اختبارات الوحدات (Unit Tests) واختبارات التكامل (Integration Tests) باستخدام Pytest للتأكد من دقة الاكتشاف وفعالية المنع.",
        "التقييم والتحسين: مراجعة أداء النظام وتحسين واجهة المستخدم بناءً على نتائج الاختبار."
    ]
    for i, step in enumerate(methodology_steps, 1):
        add_paragraph(doc, f"{i}. {step}")
    doc.add_paragraph()

    # 9. Timetable
    add_heading(doc, "9. الجدول الزمني للمشروع (Timetable for the Project)", level=2)
    add_paragraph(doc, "تم تقسيم العمل على المشروع إلى مراحل زمنية محددة لضمان الإنجاز في الوقت المحدد:")
    timetable = [
        "الشهر الأول: البحث المرجعي، تحليل المشكلة، وتحديد المتطلبات الفنية والوظيفية.",
        "الشهر الثاني: تصميم بنية النظام، وتصميم واجهات المستخدم، وإعداد بيئة التطوير.",
        "الشهر الثالث والرابع: مرحلة التطوير والبرمجة (Backend & Frontend integration)، ودمج مكتبات Presidio.",
        "الشهر الخامس: إجراء الاختبارات الشاملة للنظام، تصحيح الأخطاء، وتحسين الأداء.",
        "الشهر السادس: كتابة التقرير النهائي، إعداد العرض التقديمي، وتسليم المشروع."
    ]
    for item in timetable:
        add_bullet_point(doc, item)
    doc.add_paragraph()

    # 10. Gantt Chart
    add_heading(doc, "10. مخطط جانت للمشروع (Gantt Chart for the Project)", level=2)
    add_paragraph(doc, "يوضح مخطط جانت التوزيع الزمني للمهام وتداخلها، حيث تم تنفيذ المهام الرئيسية بالتسلسل التالي مع وجود تداخل في بعض المراحل لضمان الكفاءة:")
    gantt_desc = [
        "مرحلة التحليل: استغرقت الأسابيع الأربعة الأولى.",
        "مرحلة التصميم: بدأت في الأسبوع الخامس واستمرت لمدة 3 أسابيع.",
        "مرحلة التنفيذ: المرحلة الأطول، حيث بدأت بالتوازي مع نهاية التصميم واستمرت لمدة 8 أسابيع، وشملت بناء الـ API والواجهات.",
        "مرحلة الاختبار: خُصصت لها 4 أسابيع للتأكد من خلو النظام من الثغرات.",
        "مرحلة التوثيق: عملية مستمرة بدأت مع المشروع وتم تكثيفها في الأسابيع الأخيرة قبل التسليم."
    ]
    for item in gantt_desc:
        add_bullet_point(doc, item)
    doc.add_paragraph()

    # 11. Organization
    add_heading(doc, "11. تنظيم التقرير (Report Organization)", level=2)
    add_paragraph(doc, "تم تنظيم هذا التقرير في خمسة فصول رئيسية تغطي كافة جوانب المشروع:")
    organization = [
        "الفصل الثاني (الإطار النظري): يستعرض المفاهيم الأساسية لأمن المعلومات، أنظمة DLP، وتقنيات معالجة اللغات الطبيعية، بالإضافة إلى الدراسات السابقة.",
        "الفصل الثالث (تحليل وتصميم النظام): يشرح المتطلبات الوظيفية وغير الوظيفية، ويعرض مخططات النظام (UML Diagrams) وتصميم قاعدة البيانات.",
        "الفصل الرابع (التنفيذ والاختبار): يوضح تفاصيل الكود البرمجي، الأدوات المستخدمة في التطوير، وسيناريوهات الاختبار ونتائجها.",
        "الفصل الخامس (النتائج والتوصيات): يلخص ما تم إنجازه من أهداف، ويقدم توصيات لتطوير المشروع مستقبلاً."
    ]
    for item in organization:
        add_bullet_point(doc, item)

    # Save document
    output_path = 'Graduation_Project_Chapter1_v2.docx'
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    create_document()

