from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

from docx.oxml import OxmlElement

def set_rtl(paragraph):
    """Set paragraph direction to RTL for Arabic text."""
    pPr = paragraph._p.get_or_add_pPr()
    # Create bidi element manually
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    # Align right by default for RTL
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_heading_rtl(document, text, level=1):
    """Add an RTL heading."""
    heading = document.add_heading(text, level=level)
    set_rtl(heading)
    return heading

def add_paragraph_rtl(document, text, style=None):
    """Add an RTL paragraph."""
    paragraph = document.add_paragraph(text, style=style)
    set_rtl(paragraph)
    return paragraph

def create_document():
    document = Document()

    # --- Style Setup ---
    # Normal Text Style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'  # Standard academic font
    font.size = Pt(14)
    
    # Configure Heading Styles
    for i in range(1, 4):
        h_style = document.styles[f'Heading {i}']
        h_font = h_style.font
        h_font.name = 'Times New Roman' # Or a specific Arabic font if available, e.g., 'Simplified Arabic'
        h_font.size = Pt(16 + (4-i)*2) # H1=22, H2=20, H3=18
        h_font.bold = True
        h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # --- I. الإهداء (Dedication) ---
    document.add_section()
    add_heading_rtl(document, "الإهداء (Dedication)", level=1)
    add_paragraph_rtl(document, "\nإلى من أضاءوا لي الطريق...")
    add_paragraph_rtl(document, "إلى والديّ العزيزين، مصدر قوتي وإلهامي.")
    add_paragraph_rtl(document, "إلى أساتذتي الأفاضل الذين لم يبخلوا عليّ بعلمهم.")
    add_paragraph_rtl(document, "إلى زملائي وأصدقائي الذين شاركوني هذه الرحلة.")
    add_paragraph_rtl(document, "أهدي هذا العمل المتواضع، راجياً أن يكون خطوة نافعة في مسيرة العلم.")
    document.add_page_break()

    # --- II. التفويض (Authorization) ---
    add_heading_rtl(document, "التفويض (Authorization)", level=1)
    add_paragraph_rtl(document, "\nأفوض أنا الطالب/الطالبة (الاسم) جامعة الرازي باستخدام هذا المشروع للأغراض الأكاديمية والبحثية، مع الاحتفاظ بحقوق الملكية الفكرية للأفكار الأصيلة الواردة فيه.")
    add_paragraph_rtl(document, "\nالتوقيع: ____________________")
    add_paragraph_rtl(document, "التاريخ: ____________________")
    document.add_page_break()

    # --- III. شكر وتقدير (Acknowledgment) ---
    add_heading_rtl(document, "شكر وتقدير (Acknowledgment)", level=1)
    add_paragraph_rtl(document, "\nالحمد لله رب العالمين، والصلاة والسلام على أشرف الأنبياء والمرسلين.")
    add_paragraph_rtl(document, "أتقدم بجزيل الشكر وعظيم الامتنان إلى المشرف على هذا المشروع، الدكتور الفاضل (اسم المشرف)، على توجيهاته السديدة ودعمه المستمر.")
    add_paragraph_rtl(document, "كما أشكر كلية الحاسوب وتكنولوجيا المعلومات بجامعة الرازي على توفير البيئة الأكاديمية المحفزة.")
    add_paragraph_rtl(document, "والشكر موصول لكل من ساهم، ولو بكلمة طيبة، في إنجاز هذا العمل.")
    document.add_page_break()

    # --- IV. لجنة المناقشة (Examiner Committee) ---
    add_heading_rtl(document, "لجنة المناقشة (Examiner Committee)", level=1)
    add_paragraph_rtl(document, "\nتمت مناقشة هذا المشروع وإجازته من قبل اللجنة المكونة من:")
    add_paragraph_rtl(document, "\n1. الدكتور/ة: ____________________ (رئيساً ومشرفاً)")
    add_paragraph_rtl(document, "2. الدكتور/ة: ____________________ (مناقشاً داخلياً)")
    add_paragraph_rtl(document, "3. الدكتور/ة: ____________________ (مناقشاً خارجياً)")
    add_paragraph_rtl(document, "\nتاريخ المناقشة: __/__/2024")
    document.add_page_break()

    # --- V. الملخص (Abstract) ---
    add_heading_rtl(document, "الملخص (Abstract)", level=1)
    abstract_text = (
        "يهدف هذا المشروع إلى بناء نظام متكامل لحماية البيانات (Integrated Data Protection System) "
        "يجمع بين تقنيات منع فقدان البيانات (DLP) وتحليل النصوص باستخدام الذكاء الاصطناعي. "
        "يعالج المشروع مشكلة تسريب البيانات الحساسة (مثل أرقام الهواتف، والبيانات الشخصية PII) داخل المؤسسات، "
        "حيث يعتمد على أدوات مفتوحة المصدر قوية مثل Microsoft Presidio للتحليل الدلالي و MyDLP للمراقبة. "
        "يوفر النظام واجهة ويب تفاعلية لإدارة السياسات الأمنية، ومراقبة حركة البيانات في الوقت الفعلي، "
        "وإصدار تنبيهات فورية عند اكتشاف أي انتهاك. "
        "تم تطوير النظام باستخدام لغة Python وإطار عمل FastAPI، مع قاعدة بيانات PostgreSQL لتخزين السجلات والسياسات، "
        "مما يجعله حلاً قابلاً للتوسع وفعالاً من حيث التكلفة للمؤسسات الصغيرة والمتوسطة."
    )
    add_paragraph_rtl(document, abstract_text)
    document.add_page_break()

    # --- VI. فهرس المحتويات (Table of Contents) ---
    add_heading_rtl(document, "فهرس المحتويات (Table of Contents)", level=1)
    add_paragraph_rtl(document, "[يتم توليده تلقائياً في Word]")
    document.add_page_break()

    # --- VII. قائمة الاشكال (List of Figures) ---
    add_heading_rtl(document, "قائمة الأشكال (List of Figures)", level=1)
    add_paragraph_rtl(document, "[قائمة الأشكال]")
    document.add_page_break()

    # --- VIII. قائمة الجداول (List of Tables) ---
    add_heading_rtl(document, "قائمة الجداول (List of Tables)", level=1)
    add_paragraph_rtl(document, "[قائمة الجداول]")
    document.add_page_break()

    # --- IX. قائمة الاختصارات (List of Abbreviations) ---
    add_heading_rtl(document, "قائمة الاختصارات (List of Abbreviations)", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[1].text = 'الاختصار (Abbreviation)'
    hdr_cells[0].text = 'المعنى (Meaning)'
    
    abbreviations = [
        ("DLP", "Data Loss Prevention (منع فقدان البيانات)"),
        ("PII", "Personally Identifiable Information (معلومات تحديد الهوية الشخصية)"),
        ("GDPR", "General Data Protection Regulation (اللائحة العامة لحماية البيانات)"),
        ("HIPAA", "Health Insurance Portability and Accountability Act"),
        ("API", "Application Programming Interface"),
        ("NLP", "Natural Language Processing (معالجة اللغات الطبيعية)"),
        ("UI", "User Interface (واجهة المستخدم)")
    ]
    
    for abbr, meaning in abbreviations:
        row_cells = table.add_row().cells
        row_cells[1].text = abbr
        row_cells[0].text = meaning
        # Set RTL for table cells (basic implementation)
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                set_rtl(paragraph)

    document.add_page_break()

    # ================= الفصل الأول (Chapter One) =================
    add_heading_rtl(document, "الفصل الأول: الإطار العام للمشروع", level=1)

    # 1) مقدمة (Introduction)
    add_heading_rtl(document, "1. مقدمة (Introduction)", level=2)
    intro_text = (
        "في عصر المعلومات الرقمية، أصبحت البيانات هي الأصول الأكثر قيمة للمؤسسات. "
        "ومع تزايد حجم البيانات وتنوعها، ازدادت مخاطر تسريبها سواء عن طريق الخطأ أو الهجمات المتعمدة. "
        "تعد أنظمة منع فقدان البيانات (Data Loss Prevention - DLP) خط الدفاع الأول لحماية هذه الأصول. "
        "يقدم هذا المشروع حلاً متكاملاً يعتمد على دمج تقنيات التعلم الآلي (Machine Learning) مع القواعد التقليدية "
        "لتوفير حماية ذكية ودقيقة للبيانات الحساسة."
    )
    add_paragraph_rtl(document, intro_text)

    # 2) تعريف المشروع (Definition of the project)
    add_heading_rtl(document, "2. تعريف المشروع (Definition of the project)", level=2)
    def_text = (
        "المشروع هو نظام برمجي متكامل (Integrated System) يهدف إلى مراقبة وحماية البيانات الحساسة داخل الشبكة المؤسسية. "
        "يعتمد النظام على محرك تحليل النصوص (Microsoft Presidio) لتحديد البيانات الشخصية (PII) بدقة عالية، "
        "ويتكامل مع أدوات المراقبة (MyDLP) لمنع انتقال هذه البيانات عبر القنوات غير المصرح بها. "
        "يتميز النظام بوجود واجهة إدارة مركزية (Dashboard) تسمح للمسؤولين بوضع سياسات أمنية مرنة ومراقبة التنبيهات في الوقت الفعلي."
    )
    add_paragraph_rtl(document, def_text)

    # 3) تحديد المشكلة (Problem Statement)
    add_heading_rtl(document, "3. تحديد المشكلة (Problem Statement)", level=2)
    prob_text = (
        "تواجه المؤسسات اليوم تحديات كبيرة في حماية بياناتها، وتتمثل المشكلة الرئيسية في النقاط التالية:\n"
        "1. **قصور الأنظمة التقليدية:** تعتمد معظم أنظمة الـ DLP الحالية على الكلمات المفتاحية (Keywords) والتعابير النمطية (Regex) فقط، مما يؤدي إلى نسبة عالية من الإنذارات الخاطئة (False Positives).\n"
        "2. **التكلفة العالية:** الحلول التجارية المتقدمة باهظة الثمن ومعقدة في الإعداد.\n"
        "3. **صعوبة الإدارة:** الافتقار إلى واجهة مركزية سهلة الاستخدام لإدارة السياسات والتقارير.\n"
        "4. **البيانات غير المهيكلة:** صعوبة اكتشاف البيانات الحساسة داخل النصوص غير المهيكلة (مثل رسائل البريد الإلكتروني والملفات النصية)."
    )
    add_paragraph_rtl(document, prob_text)

    # 4) اهداف المشروع (Project Objectives)
    add_heading_rtl(document, "4. أهداف المشروع (Project Objectives)", level=2)
    objectives = (
        "يهدف المشروع إلى تحقيق الأهداف التالية:\n"
        "1. تطوير نظام قادر على اكتشاف البيانات الحساسة (مثل أرقام الهواتف، البطاقات الائتمانية، الأسماء) بدقة عالية باستخدام تقنيات NLP.\n"
        "2. بناء آلية لمنع تسريب البيانات (Prevention Mechanism) تعمل في الوقت الفعلي.\n"
        "3. توفير واجهة مستخدم (Dashboard) سهلة لإنشاء وإدارة السياسات الأمنية (Policies Management).\n"
        "4. دعم الامتثال للمعايير العالمية (مثل GDPR) من خلال تقارير وتدقيق شامل.\n"
        "5. تقديم حل مفتوح المصدر وقليل التكلفة يناسب المؤسسات المتوسطة والصغيرة."
    )
    add_paragraph_rtl(document, objectives)

    # 5) اهمية المشروع (Project Importance)
    add_heading_rtl(document, "5. أهمية المشروع (Project Importance)", level=2)
    importance = (
        "تكمن أهمية المشروع في النقاط التالية:\n"
        "1. **حماية السمعة:** منع تسريب بيانات العملاء يحمي المؤسسة من الخسائر القانونية وتضرر السمعة.\n"
        "2. **الذكاء في الحماية:** الانتقال من الحماية التقليدية الجامدة إلى الحماية الذكية القائمة على السياق.\n"
        "3. **تقليل التكاليف:** توفير بديل فعال ومجاني للأنظمة التجارية المكلفة.\n"
        "4. **تعزيز الأمن السيبراني:** سد الثغرات الأمنية المتعلقة بالعنصر البشري (Human Factor) الذي يعد أضعف حلقة في أمن المعلومات."
    )
    add_paragraph_rtl(document, importance)

    # 6) نطاق وحدود المشروع (Project Scope and Limitations)
    add_heading_rtl(document, "6. نطاق وحدود المشروع (Project Scope and Limitations)", level=2)
    scope = (
        "**النطاق (Scope):**\n"
        "- يركز المشروع على حماية البيانات النصية (Text Data).\n"
        "- يدعم اللغة الإنجليزية والعربية (بشكل أساسي).\n"
        "- يغطي قنوات التسريب عبر الويب (HTTP) ومحاكاة للبريد الإلكتروني.\n"
        "\n**الحدود (Limitations):**\n"
        "- قد لا يدعم تحليل الصور (OCR) في الإصدار الحالي.\n"
        "- يعتمد على دقة النماذج اللغوية المستخدمة، والتي قد تتأثر بالنصوص المشوشة.\n"
        "- النظام مصمم للعمل في بيئة خادم (Server-Side) ولا يشمل حماية نقاط النهاية (Endpoint Agents) بشكل كامل."
    )
    add_paragraph_rtl(document, scope)

    # 7) الادوات المستخدمة في المشروع (Tools used in the project)
    add_heading_rtl(document, "7. الأدوات المستخدمة في المشروع (Tools used in the project)", level=2)
    tools = (
        "تم استخدام مجموعة من الأدوات والتقنيات الحديثة لبناء النظام:\n"
        "1. **لغة البرمجة:** Python 3.8+ (للمرونة والدعم القوي لمكتبات الذكاء الاصطناعي).\n"
        "2. **إطار العمل الخلفي:** FastAPI (لبناء API سريع وعالي الأداء).\n"
        "3. **محرك التحليل:** Microsoft Presidio (لتحليل النصوص واكتشاف الكيانات).\n"
        "4. **محرك DLP:** MyDLP CE (كنواة لمنع التسريب).\n"
        "5. **قاعدة البيانات:** PostgreSQL (لتخزين السياسات والسجلات).\n"
        "6. **الواجهة الأمامية:** HTML5, CSS3, JavaScript (لبناء لوحة التحكم).\n"
        "7. **البيئة:** Docker & Docker Compose (لتسهيل النشر والتشغيل)."
    )
    add_paragraph_rtl(document, tools)

    # 8) المنهجية المستخدمة لعمل المشروع (Research Methodology)
    add_heading_rtl(document, "8. المنهجية المستخدمة لعمل المشروع (Research Methodology)", level=2)
    methodology = (
        "تم اتباع منهجية التطوير الرشيقة (Agile Methodology) في تنفيذ المشروع، وتحديداً إطار عمل Scrum المصغر. "
        "تم تقسيم العمل إلى مراحل (Sprints):\n"
        "1. **مرحلة التحليل:** دراسة المتطلبات وتحديد الأدوات المناسبة.\n"
        "2. **مرحلة التصميم:** تصميم هيكلية النظام (System Architecture) وقاعدة البيانات.\n"
        "3. **مرحلة التطوير:** برمجة الخدمات الأساسية (Backend) والواجهة (Frontend).\n"
        "4. **مرحلة التكامل والاختبار:** ربط المكونات وإجراء اختبارات الأداء والأمان.\n"
        "تسمح هذه المنهجية بالتطوير المستمر والاستجابة للتغييرات بسرعة."
    )
    add_paragraph_rtl(document, methodology)

    # 9) الجدول الزمني للمشروع (Timetable for the project)
    add_heading_rtl(document, "9. الجدول الزمني للمشروع (Timetable for the project)", level=2)
    add_paragraph_rtl(document, "[يمكن إدراج جدول هنا يوضح المراحل الزمنية، مثلاً: الأسبوع 1-2: التحليل، الأسبوع 3-6: التصميم والتطوير، الخ.]")

    # 10) مخطط جانت للمشروع (Gantt chart for the project)
    add_heading_rtl(document, "10. مخطط جانت للمشروع (Gantt chart for the project)", level=2)
    add_paragraph_rtl(document, "[صورة مخطط جانت توضع هنا لتمثيل الجدول الزمني بصرياً]")

    # 11) تنظيم التقرير (Report Organization)
    add_heading_rtl(document, "11. تنظيم التقرير (Report Organization)", level=2)
    org_text = (
        "يتكون هذا التقرير من الفصول التالية:\n"
        "- **الفصل الأول:** الإطار العام للمشروع (المقدمة، المشكلة، الأهداف).\n"
        "- **الفصل الثاني:** الإطار النظري والدراسات السابقة (شرح مفصل لتقنيات DLP و NLP).\n"
        "- **الفصل الثالث:** تحليل وتصميم النظام (Diagrams, Architecture).\n"
        "- **الفصل الرابع:** التنفيذ والنتائج (Implementation & Results).\n"
        "- **الفصل الخامس:** الخاتمة والتوصيات."
    )
    add_paragraph_rtl(document, org_text)

    # Save the document
    output_path = "Graduation_Project_Analysis.docx"
    document.save(output_path)
    print(f"Successfully created document at: {output_path}")

if __name__ == "__main__":
    create_document()

